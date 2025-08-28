import os
import pandas as pd
import PyPDF2
import docx
from typing import List, Dict, Any, Optional
import streamlit as st
from datetime import datetime
import hashlib
from config import config

class DocumentProcessor:
    """Handles document ingestion, processing and cleaning"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel
        }
    
    def validate_file(self, file) -> bool:
        """Validate uploaded file"""
        if not file:
            return False
        
        # Check file size
        if file.size > config.MAX_FILE_SIZE:
            st.error(f"File size ({file.size / 1024 / 1024:.1f}MB) exceeds maximum allowed size ({config.MAX_FILE_SIZE / 1024 / 1024}MB)")
            return False
        
        # Check file extension
        file_ext = os.path.splitext(file.name)[1].lower()
        if file_ext not in config.ALLOWED_EXTENSIONS:
            st.error(f"File type {file_ext} not supported. Allowed types: {', '.join(config.ALLOWED_EXTENSIONS)}")
            return False
        
        return True
    
    def process_file(self, file) -> Optional[Dict[str, Any]]:
        """Process uploaded file and extract content"""
        if not self.validate_file(file):
            return None
        
        try:
            file_ext = os.path.splitext(file.name)[1].lower()
            processor = self.supported_types.get(file_ext)
            
            if not processor:
                st.error(f"No processor available for {file_ext}")
                return None
            
            # Generate file hash for deduplication
            file_hash = hashlib.md5(file.getvalue()).hexdigest()
            
            # Process content
            content = processor(file)
            
            return {
                'filename': file.name,
                'file_type': file_ext,
                'file_hash': file_hash,
                'file_size': file.size,
                'content': content,
                'processed_at': datetime.now().isoformat(),
                'word_count': len(content.split()) if isinstance(content, str) else 0
            }
        
        except Exception as e:
            st.error(f"Error processing {file.name}: {str(e)}")
            return None
    
    def _process_pdf(self, file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return self._clean_text(text)
        
        except Exception as e:
            raise Exception(f"Failed to process PDF: {str(e)}")
    
    def _process_docx(self, file) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
        
        except Exception as e:
            raise Exception(f"Failed to process Word document: {str(e)}")
    
    def _process_csv(self, file) -> str:
        """Process CSV file"""
        try:
            df = pd.read_csv(file)
            
            # Convert to text representation
            text = f"CSV Data Summary:\n"
            text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
            text += f"Column names: {', '.join(df.columns)}\n\n"
            
            # Add sample data (first 5 rows)
            text += "Sample data (first 5 rows):\n"
            text += df.head().to_string()
            
            # Add data types and basic statistics
            text += f"\n\nData types:\n{df.dtypes.to_string()}"
            
            if len(df.select_dtypes(include='number').columns) > 0:
                text += f"\n\nNumerical summary:\n{df.describe().to_string()}"
            
            return text
        
        except Exception as e:
            raise Exception(f"Failed to process CSV: {str(e)}")
    
    def _process_excel(self, file) -> str:
        """Process Excel file"""
        try:
            excel_file = pd.ExcelFile(file)
            text = f"Excel file with {len(excel_file.sheet_names)} sheet(s):\n"
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file, sheet_name=sheet_name)
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                text += f"Columns: {', '.join(df.columns)}\n"
                
                # Add sample data
                if not df.empty:
                    text += f"Sample data:\n{df.head(3).to_string()}\n"
            
            return text
        
        except Exception as e:
            raise Exception(f"Failed to process Excel: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters but keep basic punctuation
        import re
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def get_document_stats(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not documents:
            return {}
        
        total_files = len(documents)
        total_size = sum(doc['file_size'] for doc in documents)
        total_words = sum(doc['word_count'] for doc in documents)
        
        file_types = {}
        for doc in documents:
            file_type = doc['file_type']
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_files': total_files,
            'total_size_mb': round(total_size / 1024 / 1024, 2),
            'total_words': total_words,
            'file_types': file_types,
            'avg_words_per_doc': round(total_words / total_files) if total_files > 0 else 0
        }
